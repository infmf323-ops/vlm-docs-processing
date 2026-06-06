# -*- coding: utf-8 -*-
import json, os
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
HERE=os.path.dirname(os.path.abspath(__file__))
OUT=os.path.join(HERE,"diploma.docx")
def fix_fonts(style,name="Times New Roman"):
    rpr=style.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),name)
def run_font(run,name="Times New Roman"):
    rf=run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),name)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Times New Roman"; n.font.size=Pt(14); fix_fonts(n)
pf=n.paragraph_format; pf.line_spacing=1.5; pf.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
pf.first_line_indent=Cm(1.25); pf.space_after=Pt(0); pf.space_before=Pt(0)
for nm,sz,center in [("Heading 1",16,True),("Heading 2",14,False),("Heading 3",14,False)]:
    h=doc.styles[nm]; h.font.name="Times New Roman"; h.font.size=Pt(sz); h.font.bold=True
    h.font.color.rgb=RGBColor(0,0,0); fix_fonts(h)
    hp=h.paragraph_format; hp.line_spacing=1.5; hp.space_before=Pt(12); hp.space_after=Pt(6)
    hp.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    hp.first_line_indent=Cm(0) if center else Cm(1.25); hp.keep_with_next=True
s=doc.sections[0]; s.page_width=Mm(210); s.page_height=Mm(297)
s.left_margin=Mm(30); s.right_margin=Mm(10); s.top_margin=Mm(20); s.bottom_margin=Mm(20)
def cell_border(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        e=OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"6"); e.set(qn("w:color"),"000000"); b.append(e)
    tcPr.append(b)
def para(text,style=None,align=None,indent=None,size=14,bold=False):
    p=doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None: p.paragraph_format.alignment=align
    if indent is not None: p.paragraph_format.first_line_indent=indent
    r=p.add_run(text); r.font.size=Pt(size); r.bold=bold; run_font(r); return p
RENDERED={"n":0}
_NUM={"count":0}
def render(b):
    t=b.get("t")
    if t=="h1":
        if RENDERED["n"]>0: doc.add_page_break()
        para(b["text"],style="Heading 1")
    elif t=="h2": para(b["text"],style="Heading 2")
    elif t=="h3": para(b["text"],style="Heading 3")
    elif t=="p":
        p=doc.add_paragraph(); r=p.add_run(b["text"]); r.font.size=Pt(14); run_font(r)
    elif t=="pc": para(b["text"],align=WD_ALIGN_PARAGRAPH.CENTER,indent=Cm(0),size=b.get("size",14),bold=b.get("bold",False))
    elif t in ("cap","caption_table"): para(b["text"],align=WD_ALIGN_PARAGRAPH.CENTER if t=="cap" else WD_ALIGN_PARAGRAPH.LEFT,indent=Cm(0),size=12)
    elif t in ("bullet","num"):
        if t=="num":
            _NUM["count"]+=1; prefix=f"{_NUM['count']}. "
        else:
            prefix="\u2014 "
        p=doc.add_paragraph()
        r=p.add_run(prefix+b["text"]); r.font.size=Pt(14); run_font(r)
    elif t=="img":
        p=doc.add_paragraph(); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0)
        p.add_run().add_picture(os.path.join(HERE,b["path"]), width=Mm(b.get("w",150)))
    elif t=="toc":
        p=doc.add_paragraph(); pf=p.paragraph_format
        pf.first_line_indent=Cm(0); pf.line_spacing=1.5; pf.alignment=WD_ALIGN_PARAGRAPH.LEFT
        lvl=b.get("lvl",1); pf.left_indent=Cm(0.6*(lvl-1))
        pf.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        bold=(lvl==1)
        r=p.add_run(b["text"]); r.font.size=Pt(14); r.bold=bold; run_font(r)
        r2=p.add_run("\t"+str(b["page"])); r2.font.size=Pt(14); r2.bold=bold; run_font(r2)
    elif t=="pagebreak": doc.add_page_break()
    elif t=="table":
        rows=b["rows"]; tbl=doc.add_table(rows=len(rows),cols=len(rows[0])); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        for ri,row in enumerate(rows):
            for ci,val in enumerate(row):
                c=tbl.cell(ri,ci); c.text=""; cell_border(c)
                pp=c.paragraphs[0]; pp.paragraph_format.first_line_indent=Cm(0); pp.paragraph_format.line_spacing=1.0
                rr=pp.add_run(val); rr.font.size=Pt(12); rr.bold=(ri==0); run_font(rr)
blocks=json.load(open(os.path.join(HERE,"content.json"),encoding="utf-8"))
_prev=None
for b in blocks:
    if b.get("t")=="num" and _prev!="num": _NUM["count"]=0
    render(b); RENDERED["n"]+=1; _prev=b.get("t")
doc.save(OUT); print(f"blocks={len(blocks)} saved {OUT}")
